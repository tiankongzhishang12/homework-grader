from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


OUTPUT_PATH = Path("software-project-practicum") / "长文档切片语义保持研究开题沟通稿.docx"


def set_document_style(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)

    for style_name in ("Title", "Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def add_title(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)

    sub = document.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(subtitle)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)


def add_bullets(document: Document, items: Iterable[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.line_spacing = 1.5
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)


def build_document() -> Document:
    document = Document()
    set_document_style(document)
    add_title(
        document,
        "长文档切片后跨片语义关系保持与评分一致性研究",
        "硕士阶段研究方向开题沟通稿",
    )

    document.add_heading("一、研究问题界定", level=1)
    add_paragraph(
        document,
        "本研究聚焦于长篇学生实训报告自动评分场景中的一个核心问题：当报告因上下文长度限制被切分为多个片段后，如何尽量减少跨片语义关系损失，并保持最终评分结果的一致性与稳定性。"
        "具体到软件项目基础实训报告，问题主要体现为需求规格说明、概要设计说明和详细设计说明往往跨章节、跨页分散呈现，需求点、模块设计、接口设计、数据库设计之间具有强依赖关系，若切片方式不合理或切片后缺少关系恢复机制，容易出现语义断链、证据遗漏和评分漂移。"
    )
    add_paragraph(
        document,
        "因此，本研究不是单纯讨论如何把报告切得更小，而是研究如何在切片之后保住“需求—概要—详细设计”的追踪链，确保系统既能处理长文档，又能给出证据充分、逻辑稳定的评分结果。"
    )

    document.add_heading("二、研究背景与意义", level=1)
    add_paragraph(
        document,
        "随着大语言模型在教育场景中的应用不断深入，自动批改已从短文本评分逐步扩展到课程报告、设计文档、实验报告等复杂长文档任务。与普通问答不同，实训报告评分具有明显的证据驱动特征：评分不是基于局部语言质量，而是基于跨章节的一致性、完整性和可追踪性。"
    )
    add_paragraph(
        document,
        "现有方法在长文档场景下面临两个突出挑战。其一，整文输入虽然信息完整，但推理成本高、上下文资源消耗大，且在批量评分时不具备良好的可扩展性。其二，固定窗口切片虽然能缓解上下文压力，但容易割裂原始语义结构，尤其会破坏需求与设计实现之间的对应关系，导致评分依据不稳定。"
    )
    add_paragraph(
        document,
        "因此，研究长文档切片后的低损语义保持与一致性评分方法，既具有明确的学术研究价值，也具有直接的工程应用意义。对于教育智能评测领域，该研究有助于提升复杂作业自动评分的可信度；对于长文档理解领域，该研究能够为跨片关系保持和证据聚合提供面向应用的新思路。"
    )

    document.add_heading("三、研究目标", level=1)
    add_bullets(
        document,
        [
            "设计一种面向实训报告的结构感知切片方法，使切片过程尽量贴合章节、用例、模块、页面、接口和数据库表等语义单元。",
            "构建跨片关系保持机制，将被切开的需求点、设计模块、实现对象重新连接为可追踪的语义链。",
            "提出一种面向自动评分的一致性聚合方法，降低局部切片误差对最终总分的放大效应。",
            "在真实实训报告数据上验证所提方法在语义保持、评分稳定性、证据完整性和成本效率上的有效性。",
        ],
    )

    document.add_heading("四、核心研究内容", level=1)
    document.add_heading("4.1 结构感知切片方法", level=2)
    add_paragraph(
        document,
        "研究如何从报告的自然结构出发进行切片，而不是采用简单的固定长度窗口。切片时将综合利用文档角色信息（需求、概要、详细设计）、章节标题、主题标签、关键实体锚点和评分维度信息，将文档组织成更符合评分语义的片段。"
    )
    document.add_heading("4.2 跨片语义关系保持机制", level=2)
    add_paragraph(
        document,
        "研究如何在切片之后保留和恢复跨片关系。考虑引入统一的语义标识体系，例如需求 ID、模块 ID、页面 ID、接口 ID、数据表 ID 等，并构建需求到设计再到实现的追踪图谱，使系统能够跨片定位相关证据。"
    )
    document.add_heading("4.3 一致性评分与证据聚合", level=2)
    add_paragraph(
        document,
        "研究如何将多个片段级分析结果汇聚为稳定的总评分。拟采用“先抽取证据，再聚合判断，最后统一评分”的流程，避免每个局部代理直接给出整体分数，从而降低局部偏差和评分漂移。"
    )
    document.add_heading("4.4 原型系统实现与实验验证", level=2)
    add_paragraph(
        document,
        "基于现有实训报告评分系统实现原型，将文档预处理、IR 构建、切片生成、跨片关系恢复、证据聚合和最终评分组织为完整流程，并通过对比实验验证方法有效性。"
    )

    document.add_heading("五、拟解决的关键难点", level=1)
    add_bullets(
        document,
        [
            "如何避免固定窗口切片造成的需求链断裂与图文分离。",
            "如何定义可复用、可扩展的跨片语义锚点与关系表示格式。",
            "如何处理多个局部分析结果之间的冲突、冗余和置信度差异。",
            "如何设计既能体现学术价值又便于量化验证的实验指标体系。",
        ],
    )

    document.add_heading("六、创新点凝练", level=1)
    add_paragraph(
        document,
        "从硕士论文定位看，本研究的创新点不应停留在“使用多智能体”这一表层，而应体现在方法设计和效果验证上。拟凝练为以下三个方面："
    )
    add_bullets(
        document,
        [
            "提出面向实训报告评分的结构感知切片策略，强调按语义单元而非按物理页或固定 token 窗口进行切片。",
            "提出跨片关系保持与恢复机制，通过统一语义锚点和追踪图谱重建需求—设计—实现之间的语义链。",
            "提出基于证据聚合的评分一致性优化方法，以缓解长文档切片带来的局部误判和总分波动问题。",
        ],
    )

    document.add_heading("七、技术路线", level=1)
    add_bullets(
        document,
        [
            "文档预处理：抽取章节、标题、图示说明、角色信息和主题标签，形成结构化 IR。",
            "语义切片：结合角色、章节和主题对文档进行结构化切分，并保留必要的重叠上下文。",
            "关系建模：抽取需求点、模块、页面、接口、数据表、方法流程等语义节点，构建跨片追踪关系。",
            "局部分析：对各切片执行需求抽取、覆盖判断、落地检查和证据提取。",
            "结果聚合：对局部证据进行冲突消解、置信度融合和全局仲裁，输出最终评分与评语。",
            "实验验证：与整文评分、固定窗口切片评分等基线方法进行对比。",
        ],
    )

    document.add_heading("八、实验设计初步方案", level=1)
    document.add_heading("8.1 对比方法", level=2)
    add_bullets(
        document,
        [
            "整文直接评分方法。",
            "固定窗口切片评分方法。",
            "结构切片但不做跨片关系恢复的方法。",
            "本文提出的结构切片 + 跨片关系保持 + 一致性聚合方法。",
        ],
    )
    document.add_heading("8.2 评价指标", level=2)
    add_bullets(
        document,
        [
            "评分一致性：与教师评分的一致性、多次运行结果的一致性。",
            "语义保持度：需求点在后续设计实现中的追踪成功率。",
            "证据完整性：评分结论是否具备充分证据链支撑。",
            "成本效率：token 消耗、运行时间、单份报告处理开销。",
        ],
    )
    document.add_heading("8.3 数据来源", level=2)
    add_paragraph(
        document,
        "实验数据拟来源于软件项目基础实训报告，包括需求规格说明书、概要设计说明书和详细设计说明书三类材料。可在现有评分系统基础上整理真实样本数据，并进一步构建少量教师标注集作为评价参考。"
    )

    document.add_heading("九、可行性分析", level=1)
    add_paragraph(
        document,
        "从问题来源看，该研究基于真实教学任务，问题明确且需求稳定；从技术条件看，现有系统已具备文档预处理、结构化 IR 构建、评分 rubric 和批量评分流程，可作为本研究的实验平台；从研究边界看，本研究主要依托现有大模型开展任务编排、结构化切片、关系建模和结果聚合，不依赖大规模底层模型训练，符合计算机专业硕士阶段的时间与资源条件。"
    )
    add_paragraph(
        document,
        "总体判断，该方向具有较高可行性。其优势在于问题真实、方法路径清晰、系统基础较好、实验可落地；风险主要在于数据标注工作量和评价指标设计难度，但这两点均可通过缩小样本规模、优先构建高质量验证集来控制。"
    )

    document.add_heading("十、预期成果", level=1)
    add_bullets(
        document,
        [
            "一套面向长篇实训报告评分的低损语义切片与跨片关系保持方法。",
            "一个可运行的原型系统，用于长文档自动评分实验。",
            "一组对比实验结果，用于证明方法在评分一致性和语义保持方面的有效性。",
            "一篇围绕该问题展开的硕士学位论文。",
        ],
    )

    document.add_heading("十一、建议论文题目", level=1)
    add_bullets(
        document,
        [
            "面向长篇实训报告自动评分的低损语义切片与跨片关系保持方法研究",
            "长文档切片场景下跨片语义关系保持与一致性评分方法研究",
            "面向教育智能评测的长文档结构切片与证据聚合方法研究",
        ],
    )

    document.add_heading("十二、结论", level=1)
    add_paragraph(
        document,
        "综合分析可知，以“长文档被切片后，如何尽量减少跨片语义关系损失，并保持评分一致性”为核心问题开展硕士研究，具备较高的研究价值与实施可行性。该方向位于教育智能评测、长文档理解和多智能体协同的交叉地带，既有实际应用需求，也有明确的方法研究空间。若在研究中始终围绕“低损切片、关系保持、一致性评分”三条主线展开，较有希望形成一篇问题清晰、方法完整、实验扎实的硕士论文。"
    )

    return document


def main() -> None:
    document = build_document()
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH.resolve())


if __name__ == "__main__":
    main()
