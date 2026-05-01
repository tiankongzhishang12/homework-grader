from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


OUTPUT_PATH = (
    Path(__file__).resolve().parents[1]
    / "docs"
    / "作业评分系统-数据库数据分类设计稿.docx"
)


def set_default_font(document: Document) -> None:
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(11)


def add_title(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def add_paragraph(document: Document, text: str, bold: bool = False) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)


def add_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)


def add_number_heading(document: Document, title: str) -> None:
    p = document.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")


def build_document() -> Document:
    doc = Document()
    set_default_font(doc)

    add_title(doc, "作业评分系统数据库数据分类设计稿")
    add_paragraph(doc, "文档目的：整理完整成熟作业评分系统需要长期管理和追溯的数据范围，为后续关系数据库逻辑建模、物理建模与接口设计提供统一基础。")
    add_paragraph(doc, "设计原则：本稿只讨论“应该存哪些数据”，不在此阶段展开表结构、索引、分库分表或对象存储实现细节。")

    sections = [
        (
            "1. 组织与人员数据",
            "这类数据属于基础主数据，变化频率较低，但全系统的课程、作业、提交、评分、发布都会依赖它。",
            [
                "学校：学校编号、学校名称、状态。",
                "学院：学院编号、学院名称、所属学校、状态。",
                "专业：专业编号、专业名称、所属学院、状态。",
                "班级：班级编号、班级名称、年级、所属专业、状态。",
                "老师：教师编号、姓名、所属学院、联系方式、状态。",
                "学生：学号、姓名、所属专业、所属班级、状态。",
            ],
            [
                "这部分数据的关键是主数据一致性，后续所有业务数据都必须能回溯到明确的组织和人员主体。",
                "老师与学生虽然都属于人员实体，但业务角色不同，不建议在数据库里混成一张泛用户主表后靠类型硬区分核心业务字段。",
            ],
        ),
        (
            "2. 教学与开课数据",
            "这类数据用于描述“谁在什么时间教什么课”，是连接组织主数据与考核评分数据的核心桥梁。",
            [
                "课程：课程编号、课程名称、学分、课程类型、开设单位。",
                "课程开设：所属课程、学年、学期、开设名称、状态。",
                "教学班：所属课程开设、教学班名称、状态。",
                "任课老师关联：课程开设、老师、角色（主讲、辅讲、阅卷负责人等）。",
                "教学班学生关联：教学班、学生、选课或在班状态。",
            ],
            [
                "课程本身是主数据，课程开设是具体业务实例，两者不能混为一谈。",
                "任课老师和选课学生都应以关联关系存储，而不是仅靠课程表或班级表中的单字段挂接。",
            ],
        ),
        (
            "3. 考核项数据",
            "这类数据描述某门课下面有哪些平时作业、实训报告、测验、期末试卷，以及它们各自的提交和评分约束。",
            [
                "平时作业。",
                "实训报告。",
                "测验。",
                "期末试卷。",
            ],
            [
                "每个考核项至少要存：名称、所属课程开设、类型、总分、权重、截止时间、提交形式、评分模式。",
                "必须明确区分：这是平时作业还是期末试卷；是答题卡图片还是 Word/PDF；是按题评分还是按 rubric 评分。",
                "如果系统未来支持混合作业，也应允许一个考核项定义多资产提交规则，而不是默认只接收单一文件。",
            ],
        ),
        (
            "4. 模板与评分规则数据",
            "这部分是自动评分系统的配置核心，决定系统如何识别、如何提取、如何评分、如何生成批注结果。",
            [
                "考核模板。",
                "题目定义。",
                "评分维度定义。",
                "rubric 版本。",
                "标准答案。",
                "答题卡坐标模板。",
            ],
            [
                "对于答题卡类考核，需要存：题号、题型、满分、区域坐标、标准答案、批注位置。",
                "对于报告类考核，需要存：评分维度、每个维度满分、评分标准、扣分规则。",
                "模板与评分规则必须有版本概念，否则后续无法解释同一份作业为什么在不同时间点得到不同分数。",
            ],
        ),
        (
            "5. 学生提交数据",
            "这类数据描述学生具体交了什么、什么时候交、交了几次、当前处于什么状态，是评分链路的起点。",
            [
                "提交记录。",
                "提交时间。",
                "提交状态。",
                "提交次数。",
                "是否逾期。",
            ],
            [
                "一个提交下面还应存它的文件资产，且文件与提交记录要分离建模。",
                "文件资产可能包括：原始图片、Word 文档、PDF、压缩包、归一化图片、批注图、预览文件。",
                "从教学平台同步来的作业文件，不应只落本地文件夹，还要在数据库中登记文件与提交、学生、考核项之间的关系。",
            ],
        ),
        (
            "6. 内容提取数据",
            "这部分是机器在正式评分之前产出的中间理解结果，用于区分“识别问题”和“评分问题”。",
            [
                "OCR 结果。",
                "文档文本提取结果。",
                "题目识别结果。",
                "页面结构。",
                "证据片段。",
                "模型提取摘要。",
            ],
            [
                "这层数据建议保留完整原始 JSON 结果，同时数据库中保存提取运行记录、结果路径、模型信息、状态和版本信息。",
                "若不存这层数据，后续无法准确定位错误究竟来自图像识别、文本提取、结构理解还是评分逻辑本身。",
            ],
        ),
        (
            "7. 评分结果数据",
            "这部分是系统真正的评分产出，必须具备可追溯、可复核、可对比的能力。",
            [
                "评分运行记录。",
                "每次评分总分。",
                "每个评分项得分。",
                "每题得分或每维度得分。",
                "是否正确。",
                "是否需要复核。",
                "评分证据。",
                "评分置信度。",
            ],
            [
                "至少要区分三层结果：自动评分结果、人工复核结果、最终确认结果。",
                "不要只存最终分数，否则系统无法解释评分过程，也无法为老师提供复核和修订依据。",
                "评分结果应与提取结果分层存储，避免把“识别输出”和“评分结论”混成一个文件或一张表。",
            ],
        ),
        (
            "8. 成绩与发布数据",
            "这类数据是老师最终使用和面向学生发布的数据，应视为业务确认结果，而不是算法中间结果。",
            [
                "单次作业最终成绩。",
                "期末试卷最终成绩。",
                "课程总评。",
                "成绩发布时间。",
                "发布状态。",
                "是否已确认。",
            ],
            [
                "课程总评如果需要长期保存、导出、申诉或审计，应单独存储，不建议只在查询时临时计算。",
                "发布状态与确认状态要分开，避免出现“成绩算出来了但还没对外发布”与“成绩已发布但仍允许复核”的语义混乱。",
            ],
        ),
    ]

    for title, intro, bullets, notes in sections:
        add_number_heading(doc, title)
        add_paragraph(doc, intro)
        add_paragraph(doc, "建议存储内容：", bold=True)
        for item in bullets:
            add_bullet(doc, item)
        add_paragraph(doc, "设计说明：", bold=True)
        for note in notes:
            add_bullet(doc, note)

    add_number_heading(doc, "9. 本阶段结论")
    add_paragraph(doc, "当前数据库数据范围应至少覆盖上述 8 大类内容，并明确区分主数据、业务数据、文件资产、中间提取数据、评分结果数据、最终发布数据。")
    add_paragraph(doc, "下一步建议在本稿基础上继续输出：概念 ER 图、逻辑模型、核心表字段清单、文件存储规范和同步接口映射方案。")
    return doc


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
