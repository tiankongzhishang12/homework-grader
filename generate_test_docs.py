from docx import Document
from docx.shared import Pt
import os

def create_doc(filename, content_lines, title):
    doc = Document()
    # 添加标题
    heading = doc.add_heading(title, 0)
    heading.alignment = 1  # 居中
    
    # 添加内容
    for line in content_lines:
        if line.startswith("##"):
            doc.add_heading(line.replace("##", ""), level=2)
        elif line.startswith("###"):
            doc.add_heading(line.replace("###", ""), level=3)
        else:
            doc.add_paragraph(line)
    
    # 保存
    doc.save(filename)
    print(f"✅ 已生成: {filename}")

# 定义路径
output_dir = "workspace/raw"
os.makedirs(output_dir, exist_ok=True)

# --- 文档 1: 高分作业 (张三) ---
# 特点：引用 GB 标准，结构完整 (问题识别，影响评估，对策建议)，建议具体
zhang_content = [
    "学生姓名：张三",
    "学号：2025001",
    "",
    "## 1. 问题识别",
    "本项目针对某化工厂周边的土壤重金属污染问题进行识别。经初步调查，主要污染物为铅 (Pb)、镉 (Cd) 和汞 (Hg)。污染源明确为工厂长期违规排放废水。",
    "",
    "## 2. 影响评估",
    "依据《土壤环境质量 农用地土壤污染风险管控标准 (试行)》(GB 15618-2018) 进行风险评估。",
    "监测数据显示，采样点 S1 的镉含量为 1.5 mg/kg，超过 GB 15618-2018 规定的风险筛选值 (0.3 mg/kg) 达 5 倍。",
    "同时参考《建设用地土壤污染风险管控和修复监测技术导则》(HJ 25.1-2019)，确认该区域存在较高的健康风险，需立即采取管控措施。",
    "定量评估显示，若不及时修复，未来 5 年内周边农作物超标概率将达到 80%。",
    "",
    "## 3. 对策建议",
    "基于上述评估，提出以下分层级建议：",
    "### 3.1 短期措施 (1 个月内)",
    "- 立即设立警戒线，禁止周边农户耕种和采摘。",
    "- 对高浓度污染点 (S1, S2) 进行覆盖固化，防止扬尘扩散。",
    "",
    "### 3.2 中期措施 (6 个月内)",
    "- 启动土壤淋洗修复工程，预计去除率可达 70%。",
    "- 建立地下水监测井，每月监测一次重金属指标。",
    "",
    "### 3.3 长期措施 (1-3 年)",
    "- 调整土地利用规划，将该地块调整为工业用地或生态林地。",
    "- 引入植物修复技术，种植超富集植物持续净化土壤。",
    "",
    "## 4. 结论",
    "本报告严格遵循 GB 15618-2018 及 HJ 25.1-2019 标准，数据详实，建议具有高度可操作性。"
]

# --- 文档 2: 低分作业 (李四) ---
# 特点：无标准引用，结构缺失 (无影响评估章节)，建议空泛
lisi_content = [
    "学生姓名：李四",
    "学号：2025002",
    "",
    "## 1. 问题识别",
    "我觉得这个工厂旁边可能有污染。听说有人看到水变黑了，可能有一些重金属，比如铅什么的。但是我没有具体数据，只是大概看了一下。",
    "",
    "## 2. 对策建议",
    "我觉得应该让工厂停止排放。然后政府应该管一管。",
    "建议大家不要在这里种菜了，对身体不好。",
    "希望能尽快解决这个问题，保护环境很重要。",
    "",
    "## 3. 总结",
    "总之，污染很严重，需要大家共同努力。报告写完了。"
]

# 生成文件
create_doc(os.path.join(output_dir, "student_001_zhang.docx"), zhang_content, "环境评估报告 - 张三")
create_doc(os.path.join(output_dir, "student_002_lisi.docx"), lisi_content, "环境评估报告 - 李四")

print("\n🎉 测试文档生成完毕！请重新运行预处理和评分流程。")
